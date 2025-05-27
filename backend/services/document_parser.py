from abc import ABC, abstractmethod
import PyPDF2
from docx import Document
import pytesseract
from PIL import Image
from typing import Dict, Any, List
import logging

class DocumentParser(ABC):
    @abstractmethod
    def parse(self, file_path: str) -> Dict[str, Any]:
        pass

class PDFParser(DocumentParser):
    def parse(self, file_path: str) -> Dict[str, Any]:
        try:
            text = ""
            metadata = {}
            
            with open(file_path, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                metadata = {
                    'pages': len(reader.pages),
                    'title': reader.metadata.get('/Title', 'Unknown') if reader.metadata else 'Unknown'
                }
                
                for page_num, page in enumerate(reader.pages, 1):
                    page_text = page.extract_text()
                    if page_text.strip():  # Skip empty pages
                        text += f"\n--- Page {page_num} ---\n{page_text}"
            
            return {
                'text': text.strip(),
                'metadata': metadata,
                'type': 'pdf',
                'sections': self._extract_sections(text)
            }
        except Exception as e:
            logging.error(f"PDF parsing error: {e}")
            raise

    def _extract_sections(self, text: str) -> List[Dict[str, str]]:
        """Extract sections based on headers and structure"""
        lines = text.split('\n')
        sections = []
        current_section = {'title': 'Introduction', 'content': ''}
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Simple header detection (can be improved with NLP)
            if (len(line) < 100 and 
                (line.isupper() or 
                 any(line.startswith(prefix) for prefix in ['Chapter', 'Section', '1.', '2.', '3.']))):
                
                if current_section['content']:
                    sections.append(current_section)
                current_section = {'title': line, 'content': ''}
            else:
                current_section['content'] += line + ' '
        
        if current_section['content']:
            sections.append(current_section)
            
        return sections

class DocxParser(DocumentParser):
    def parse(self, file_path: str) -> Dict[str, Any]:
        try:
            doc = Document(file_path)
            text = ""
            sections = []
            
            for paragraph in doc.paragraphs:
                if paragraph.style.name.startswith('Heading'):
                    sections.append({'title': paragraph.text, 'content': ''})
                else:
                    text += paragraph.text + "\n"
                    if sections:
                        sections[-1]['content'] += paragraph.text + " "
            
            return {
                'text': text.strip(),
                'metadata': {'paragraphs': len(doc.paragraphs)},
                'type': 'docx',
                'sections': sections
            }
        except Exception as e:
            logging.error(f"DOCX parsing error: {e}")
            raise

# Parser Factory
class ParserFactory:
    parsers = {
        'pdf': PDFParser,
        'docx': DocxParser
    }
    
    @classmethod
    def get_parser(cls, file_type: str) -> DocumentParser:
        parser_class = cls.parsers.get(file_type.lower())
        if not parser_class:
            raise ValueError(f"Unsupported file type: {file_type}")
        return parser_class()